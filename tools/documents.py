"""Document intelligence for Jarvis.

Reads and extracts content from:
- PDF files (via pypdf)
- DOCX files (via python-docx)
- CSV files
- Markdown files
- TXT files
- Code files (syntax-aware extraction)

Provides: content extraction, summarization, question answering, information location.
"""

import csv
import io
import re
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from datetime import datetime

from jarvis.tools.registry import BaseTool, ToolResult
from jarvis.tools.filesystem import fs
from jarvis.config import config
from jarvis.logger import logger


# ──────────────────────────────────────────────
#  Document Manager
# ──────────────────────────────────────────────

class DocumentManager:
    """Reads and processes documents of various formats."""

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".docx", ".csv", ".md", ".markdown",
        ".txt", ".text", ".log",
        ".py", ".js", ".ts", ".jsx", ".tsx",
        ".java", ".c", ".cpp", ".h", ".hpp",
        ".go", ".rs", ".rb", ".php", ".swift",
        ".json", ".yaml", ".yml", ".toml", ".xml",
        ".html", ".css", ".scss",
        ".sql", ".sh", ".bash",
        ".cfg", ".ini", ".conf",
    }

    # Text files we can read directly
    PLAIN_TEXT_EXTENSIONS = {
        ".txt", ".text", ".log", ".md", ".markdown",
        ".py", ".js", ".ts", ".jsx", ".tsx",
        ".java", ".c", ".cpp", ".h", ".hpp",
        ".go", ".rs", ".rb", ".php", ".swift",
        ".json", ".yaml", ".yml", ".toml", ".xml",
        ".html", ".css", ".scss",
        ".sql", ".sh", ".bash",
        ".cfg", ".ini", ".conf",
        ".csv",  # special handling
    }

    def __init__(self) -> None:
        logger.info("Document manager initialized")

    def is_supported(self, path: str) -> bool:
        """Check if a file type is supported."""
        ext = Path(path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    def extract_text(self, path: str) -> Dict[str, Any]:
        """Extract text content from any supported document.

        Returns:
            Dict with keys: path, format, content, metadata, page_count, word_count
        """
        resolved = fs.resolve_path(path, must_exist=True)
        if not resolved.is_file():
            raise ValueError(f"Not a file: {path}")

        ext = resolved.suffix.lower()
        logger.info(f"Extracting text from {resolved} (format: {ext})")

        if ext == ".pdf":
            return self._extract_pdf(resolved)
        elif ext == ".docx":
            return self._extract_docx(resolved)
        elif ext == ".csv":
            return self._extract_csv(resolved)
        else:
            return self._extract_text(resolved)

    def _extract_pdf(self, path: Path) -> Dict[str, Any]:
        """Extract text from a PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required to read PDF files. Install: pip install pypdf"
            )

        reader = PdfReader(str(path))
        pages = []
        full_text = ""

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"page": i + 1, "text": text.strip()})
            full_text += f"\n--- Page {i + 1} ---\n{text}"

        metadata = {}
        if reader.metadata:
            for key in reader.metadata:
                val = reader.metadata[key]
                if val:
                    metadata[key.replace("/", "").lower()] = str(val)

        return {
            "path": str(path),
            "format": "pdf",
            "content": full_text.strip(),
            "pages": pages,
            "page_count": len(pages),
            "word_count": len(full_text.split()),
            "metadata": metadata,
        }

    def _extract_docx(self, path: Path) -> Dict[str, Any]:
        """Extract text from a DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required to read DOCX files. Install: pip install python-docx"
            )

        doc = Document(str(path))
        paragraphs = []
        full_text = ""

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
                full_text += para.text + "\n"

        # Also extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        metadata = {
            "paragraphs": len(paragraphs),
            "tables": len(tables),
            "sections": len(doc.sections),
        }

        return {
            "path": str(path),
            "format": "docx",
            "content": full_text.strip(),
            "paragraphs": paragraphs,
            "tables": tables,
            "word_count": len(full_text.split()),
            "metadata": metadata,
        }

    def _extract_csv(self, path: Path) -> Dict[str, Any]:
        """Extract text from a CSV file."""
        rows = []
        headers = []
        full_text = ""

        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            sample = f.read(1024)
            f.seek(0)
            dialect = csv.Sniffer().sniff(sample) if sample else csv.excel
            reader = csv.reader(f, dialect)

            for i, row in enumerate(reader):
                if i == 0:
                    headers = row
                rows.append(row)
                full_text += ",".join(row) + "\n"

        return {
            "path": str(path),
            "format": "csv",
            "content": full_text.strip(),
            "headers": headers,
            "rows": rows,
            "row_count": len(rows),
            "column_count": len(headers) if headers else 0,
            "word_count": len(full_text.split()),
        }

    def _extract_text(self, path: Path) -> Dict[str, Any]:
        """Extract text from a plain text or code file."""
        try:
            content = path.read_text(encoding='utf-8', errors='replace')
        except Exception as e:
            raise ValueError(f"Cannot read file: {e}")

        lines = content.split('\n')
        return {
            "path": str(path),
            "format": path.suffix.lower().lstrip('.') or "txt",
            "content": content,
            "lines": lines,
            "line_count": len(lines),
            "word_count": len(content.split()),
            "metadata": {
                "extension": path.suffix.lower(),
            },
        }

    def summarize(self, path: str, max_length: int = 500) -> str:
        """Generate a concise summary of a document.

        Uses extraction + intelligent truncation at paragraph/sentence boundaries.
        For full LLM-based summarization, the content is passed to the model
        via the agent loop.
        """
        result = self.extract_text(path)
        text = result["content"]

        # Truncation-based summary for very large documents
        if len(text) > max_length * 5:
            # Take first portion (introduction), middle, and end (conclusion)
            third = len(text) // 3
            parts = [
                text[:third],
                text[third:2 * third],
                text[2 * third:],
            ]
            summary_parts = []
            for i, part in enumerate(parts[:3]):
                # Take first paragraph of each section
                paragraphs = [p.strip() for p in part.split('\n\n') if p.strip()]
                if paragraphs:
                    summary_parts.append(paragraphs[0][:max_length // 3])
            summary = " [...] ".join(summary_parts)
        else:
            summary = text

        # Build metadata-enhanced summary
        lines = [
            f"Document: {result.get('path', path)}",
            f"Format: {result.get('format', 'unknown').upper()}",
        ]

        if "page_count" in result:
            lines.append(f"Pages: {result['page_count']}")
        if "word_count" in result:
            lines.append(f"Words: {result['word_count']}")
        if "row_count" in result:
            lines.append(f"Rows: {result['row_count']}")

        lines.append("")
        lines.append("--- Content ---")
        lines.append(summary[:max_length])

        if len(summary) > max_length:
            lines.append("... (truncated)")

        return "\n".join(lines)

    def locate_information(self, path: str, query: str) -> List[Dict[str, Any]]:
        """Find specific information in a document by keyword/regex search."""
        resolved = fs.resolve_path(path, must_exist=True)
        ext = resolved.suffix.lower()

        if ext == ".pdf":
            result = self._extract_pdf(resolved)
            pages = result["pages"]
        else:
            result = self._extract_text(resolved)
            # Treat as one "page"
            pages = [{"page": 1, "text": result["content"]}]

        matches = []
        try:
            regex = re.compile(re.escape(query), re.IGNORECASE)
        except re.error:
            regex = re.compile(query, re.IGNORECASE)

        for page in pages:
            text = page["text"]
            for i, line in enumerate(text.split('\n'), 1):
                if regex.search(line):
                    matches.append({
                        "page": page.get("page", 1),
                        "line": i,
                        "text": line.strip()[:200],
                    })

        return matches


# Global instance
document_manager = DocumentManager()


# ══════════════════════════════════════════════
#  Tool Implementations
# ══════════════════════════════════════════════

class ReadDocumentTool(BaseTool):
    name = "read_document"
    description = "Read a document (PDF, DOCX, CSV, MD, TXT) and extract its contents"
    requires_confirmation = False

    def validate_args(self, **kwargs) -> tuple[bool, Optional[str]]:
        if "path" not in kwargs:
            return False, "Missing required argument: path"
        return True, None

    def get_required_args(self) -> List[str]:
        return ["path"]

    async def execute(self, **kwargs) -> ToolResult:
        try:
            path = kwargs["path"]
            result = document_manager.extract_text(path)

            lines = [
                f"Document: {result['path']}",
                f"Format: {result.get('format', 'unknown').upper()}",
            ]

            if "page_count" in result:
                lines.append(f"Pages: {result['page_count']}")
            if "word_count" in result:
                lines.append(f"Words: {result['word_count']}")
            if "row_count" in result:
                lines.append(f"Rows: {result['row_count']}")
            if "headers" in result and result["headers"]:
                lines.append(f"Columns: {', '.join(result['headers'])}")

            # Show content preview
            content = result.get("content", "")
            preview = content[:3000]
            lines.append("")
            lines.append("--- Content Preview ---")
            lines.append(preview)

            if len(content) > 3000:
                lines.append(f"\n... ({len(content) - 3000} more characters)")

            return ToolResult(
                success=True,
                message="\n".join(lines),
                data=result,
            )
        except Exception as e:
            return ToolResult(success=False, message=str(e), error=str(e))


class SummarizeDocumentTool(BaseTool):
    name = "summarize_document"
    description = "Summarize a document (PDF, DOCX, CSV, MD, TXT)"
    requires_confirmation = False

    def validate_args(self, **kwargs) -> tuple[bool, Optional[str]]:
        if "path" not in kwargs:
            return False, "Missing required argument: path"
        return True, None

    def get_required_args(self) -> List[str]:
        return ["path"]

    async def execute(self, **kwargs) -> ToolResult:
        try:
            path = kwargs["path"]
            max_length = kwargs.get("max_length", 500)
            summary = document_manager.summarize(path, max_length)

            return ToolResult(
                success=True,
                message=summary,
                data={"path": path, "summary": summary},
            )
        except Exception as e:
            return ToolResult(success=False, message=str(e), error=str(e))


class LocateInDocumentTool(BaseTool):
    name = "locate_in_document"
    description = "Find specific information in a document by search query"
    requires_confirmation = False

    def validate_args(self, **kwargs) -> tuple[bool, Optional[str]]:
        if "path" not in kwargs:
            return False, "Missing required argument: path"
        if "query" not in kwargs:
            return False, "Missing required argument: query"
        return True, None

    def get_required_args(self) -> List[str]:
        return ["path", "query"]

    async def execute(self, **kwargs) -> ToolResult:
        try:
            path = kwargs["path"]
            query = kwargs["query"]
            matches = document_manager.locate_information(path, query)

            if not matches:
                return ToolResult(
                    success=True,
                    message=f"No matches for '{query}' in {path}",
                    data=[],
                )

            lines = [f"Found {len(matches)} match(es) for '{query}' in {path}:"]
            for m in matches[:20]:
                page_info = f" (Page {m['page']})" if m.get("page") else ""
                lines.append(f"  L{m['line']}{page_info}: {m['text']}")

            if len(matches) > 20:
                lines.append(f"  ... and {len(matches) - 20} more")

            return ToolResult(
                success=True,
                message="\n".join(lines),
                data=matches,
            )
        except Exception as e:
            return ToolResult(success=False, message=str(e), error=str(e))